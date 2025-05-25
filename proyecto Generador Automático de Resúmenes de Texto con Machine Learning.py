import streamlit as st
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from transformers import T5Tokenizer, T5ForConditionalGeneration
from io import StringIO
import docx
from PIL import Image
import base64
import tempfile
from reportlab.pdfgen import canvas
import nltk
nltk.download("punkt")

# Cargar la imagen local
with open("imagen fondo proyecto.jpg", "rb") as img_file:
    img_bytes = img_file.read()
    img_base64 = base64.b64encode(img_bytes).decode()

# Estilos personalizados con CSS
# Insertar imagen de fondo usando base64
st.markdown(
    f"""
    <style>
    .stApp {{
        background-image: url('data:image/jpg;base64,{img_base64}');
        background-size: cover;
        background-position: center;
        background-repeat: no-repeat;
        color: #FFFFFF;
    }}

     label, .stSelectbox label, .stSlider label, .stFileUploader label {{
        color: white !important;
    }}

    .stButton > button {{
        background-color: #4CAF50;
        color: white;
        border-radius: 10px;
        padding: 0.5em 1em;
        font-size: 18px;
    }}

    .stTextInput > div > div > input, .stTextArea > div > textarea {{
        background-color: rgba(0, 0, 0, 0.6);
        color: #FFFFFF;
        border: 1px solid #ccc;
        border-radius: 10px;
    }}

    .css-1v0mbdj, .css-1d391kg, .st-cj {{
        background-color: rgba(0, 0, 0, 0.5);
        border-radius: 15px;
        padding: 1em;

     }}

    .css-1cpxqw2, .css-q8sbsg, .css-10trblm {{
        color: white !important;
        
    }}
    </style>
    """,
    unsafe_allow_html=True
)


# Función para leer archivo .txt
def leer_txt(archivo):
    return archivo.getvalue()

# Función para leer archivo .docx
def leer_docx(archivo):
    doc = docx.Document(archivo)
    texto = ""
    for parrafo in doc.paragraphs:
        texto += parrafo.text + "\n"
    return texto

# Función para resumen extractivo
def resumen_extractivo(texto, num_oraciones=3):
    parser = PlaintextParser.from_string(texto, Tokenizer("spanish"))
    summarizer = LsaSummarizer()
    resumen = summarizer(parser.document, num_oraciones)
    return ' '.join(str(oracion) for oracion in resumen)

# Función para resumen abstractive usando T5
def resumen_abstractive(texto, max_input_length=512, max_output_length=150):
    modelo = T5ForConditionalGeneration.from_pretrained("t5-base")
    tokenizador = T5Tokenizer.from_pretrained("t5-base")

    texto_procesado = "summarize: " + texto
    entradas = tokenizador.encode(texto_procesado, return_tensors="pt", max_length=max_input_length, truncation=True)
    salida = modelo.generate(entradas, max_length=max_output_length, min_length=120, length_penalty=2.0, num_beams=4, no_repeat_ngram_size=3, early_stopping=True)
    resumen = tokenizador.decode(salida[0], skip_special_tokens=True)

    return resumen

# Función para contar palabras
def contar_palabras(texto):
    return len(texto.split())

def guardar_resumen_en_docx(resumen):
    doc = docx.Document()
    doc.add_heading("Resumen Generado", level=1)
    doc.add_paragraph(resumen)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def guardar_resumen_en_pdf(resumen):
    buffer = BytesIO()
    c = canvas.Canvas(buffer)
    c.setFont("Helvetica", 12)
    text_object = c.beginText(40, 800)
    for linea in resumen.split('\n'):
        for sub_linea in [linea[i:i+90] for i in range(0, len(linea), 90)]:
            text_object.textLine(sub_linea)
    c.drawText(text_object)
    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer


# Interfaz de Streamlit
st.title("📝 Generador Automático de Resúmenes de Texto")
st.write("Introduce un texto largo, carga un archivo, o elige el tipo de resumen que deseas obtener:")

uploaded_file = st.file_uploader("Cargar archivo (.txt o .docx)", type=["txt", "docx"])
texto_largo = st.text_area("O introduce tu texto largo aquí:", height=300)

opcion = st.selectbox("Selecciona el tipo de resumen", ("Resumen Extractivo", "Resumen Abstractive"))
if opcion == "Resumen Extractivo":
    num_oraciones = st.slider("Número de oraciones en el resumen", min_value=1, max_value=10, value=3)
else:
    max_palabras = st.slider("Máximo de palabras en el resumen", min_value=50, max_value=5000, value=300, step=10)

if st.button("🔍 Generar Resumen"):
    if uploaded_file is not None:
        if uploaded_file.type == "text/plain":
            texto_largo = leer_txt(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            texto_largo = leer_docx(uploaded_file)
    
    if texto_largo:
        with st.spinner("Generando resumen..."):
            if opcion == "Resumen Extractivo":
                resumen = resumen_extractivo(texto_largo, num_oraciones=num_oraciones)
            else:
                resumen = resumen_abstractive(texto_largo, max_output_length=max_palabras)

        st.success("¡Resumen generado exitosamente!")
        st.subheader("📄 Resumen:")
        st.write(resumen)

        num_palabras_resumen = contar_palabras(resumen)
        st.write(f"📝 El resumen contiene {num_palabras_resumen} palabras.")

        # Descargar como .txt
        resumen_bytes = resumen.encode('utf-8')
        st.download_button(
            label="💾 Descargar como TXT",
            data=resumen_bytes,
            file_name="resumen.txt",
            mime="text/plain"
        )

        # Descargar como .docx
        docx_buffer = guardar_resumen_en_docx(resumen)
        st.download_button(
            label="📄 Descargar como DOCX",
            data=docx_buffer,
            file_name="resumen.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # Descargar como .pdf
        pdf_buffer = guardar_resumen_en_pdf(resumen)
        st.download_button(
            label="📑 Descargar como PDF",
            data=pdf_buffer,
            file_name="resumen.pdf",
            mime="application/pdf"
        )

    else:
        st.error("⚠️ Por favor, introduce o carga un texto para generar el resumen.")
